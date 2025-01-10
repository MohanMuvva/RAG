import os
import fitz  # PyMuPDF for PDF
import docx  # For DOCX files
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
import hashlib
from typing import Dict
import time
from datetime import datetime

# Define the persist directory
PERSIST_DIRECTORY = "C:\\Users\\saich\\Favorites\\Mohan\\RAG\\RAG\\DOCUMENT_RAG\\local_chroma_db"

# Debug flag
DEBUG = True  # Set to False to suppress debug output

# Ensure the directory exists
os.makedirs(PERSIST_DIRECTORY, exist_ok=True)

def debug_print(message):
    """Print debug messages if DEBUG is True."""
    if DEBUG:
        print(f"[DEBUG] {message}")

# Initialize ChromaDB Client
def initialize_chromadb():
    try:
        debug_print("Initializing ChromaDB...")
        client = chromadb.Client(Settings(
            chroma_api_impl="chromadb.api.fastapi.FastAPI",
            chroma_server_host="localhost",
            chroma_server_http_port=8000,
            persist_directory=PERSIST_DIRECTORY
        ))
        debug_print(f"ChromaDB initialized with persist directory: {PERSIST_DIRECTORY}")
        return client
    except Exception as e:
        debug_print(f"Error initializing ChromaDB: {str(e)}")
        return None

# Generate content hash
def get_content_hash(file_path: str) -> str:
    try:
        if file_path.endswith('.pdf'):
            doc = fitz.open(file_path)
            content = ' '.join(page.get_text() for page in doc)
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            content = ' '.join(paragraph.text for paragraph in doc.paragraphs)
        return hashlib.md5(content.encode()).hexdigest()
    except Exception as e:
        debug_print(f"Error generating hash for {file_path}: {str(e)}")
        return ""

# Extract text from files
def extract_text_from_file(file_path):
    try:
        if file_path.endswith(".pdf"):
            doc = fitz.open(file_path)
            text = [page.get_text() for page in doc]
            return " ".join(text)
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            text = [paragraph.text for paragraph in doc.paragraphs]
            return "\n".join(text)
    except Exception as e:
        debug_print(f"Error extracting text from {file_path}: {str(e)}")
        return ""

# Chunk text
def chunk_text(text, chunk_size=1000, overlap=200):
    try:
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap  # Overlap for context preservation
        return chunks
    except Exception as e:
        debug_print(f"Error during text chunking: {str(e)}")
        return []

# Generate embeddings
def generate_embeddings(chunks, model_name="sentence-transformers/all-MiniLM-L6-v2"):
    try:
        model = SentenceTransformer(model_name)
        embeddings = model.encode(chunks, show_progress_bar=True)
        return embeddings
    except Exception as e:
        debug_print(f"Error generating embeddings: {str(e)}")
        return []

# Retrieve existing chunks
def get_existing_chunks(file_name, collection):
    try:
        results = collection.get(where={"source_file": file_name})
        return {result["chunk_index"]: result.get("document", None) for result in results["metadatas"]}
    except Exception as e:
        debug_print(f"Error retrieving existing chunks from ChromaDB: {str(e)}")
        return {}

# Store chunks in ChromaDB
def store_in_chromadb(file_name, chunks, embeddings, start_chunk_idx, collection):
    try:
        if not collection:
            debug_print("Error: ChromaDB collection is not initialized.")
            return

        if not chunks or not embeddings:
            debug_print(f"Error: Chunks or embeddings are empty for {file_name}.")
            return
        if len(chunks) != len(embeddings):
            debug_print(f"Error: Mismatch in chunks and embeddings lengths for {file_name}.")
            debug_print(f"Chunks: {len(chunks)}, Embeddings: {len(embeddings)}")
            return

        debug_print(f"Storing {len(chunks)} chunks in ChromaDB starting at index {start_chunk_idx}...")
        for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings), start=start_chunk_idx):
            unique_id = f"{file_name}_chunk_{idx}"
            debug_print(f"Storing chunk {idx}: {chunk[:50]} with embedding length {len(embedding)}")
            collection.add(
                ids=[unique_id],
                metadatas=[{"chunk_index": idx, "source_file": file_name}],
                documents=[chunk],
                embeddings=[embedding.tolist()]
            )
        debug_print(f"Stored {len(chunks)} chunks for {file_name}.")

    except Exception as e:
        debug_print(f"Error storing chunks in ChromaDB: {str(e)}")

# Process files
def process_files(content_hashes, collection):
    folder_path = "C:\\Users\\saich\\Favorites\\Mohan\\RAG\\RAG\\DOCUMENT_RAG"
    current_files = {
        f for f in os.listdir(folder_path)
        if f.endswith((".pdf", ".docx")) and not f.startswith(('~', '.'))
    }

    for file_name in current_files:
        file_path = os.path.join(folder_path, file_name)

        if not os.access(file_path, os.R_OK):
            debug_print(f"Skipping inaccessible file: {file_name}")
            continue

        current_hash = get_content_hash(file_path)
        if current_hash and file_name in content_hashes and content_hashes[file_name] == current_hash:
            debug_print(f"No changes detected for {file_name}. Skipping.")
            continue

        debug_print(f"Processing file: {file_name}")

        file_text = extract_text_from_file(file_path)
        chunks = chunk_text(file_text)
        embeddings = generate_embeddings(chunks)

        existing_chunks = get_existing_chunks(file_name, collection)
        modified_chunks = [
            (idx, chunk, embedding)
            for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings))
            if idx not in existing_chunks or existing_chunks[idx] != chunk
        ]

        if modified_chunks:
            store_in_chromadb(
                file_name,
                [chunk for _, chunk, _ in modified_chunks],
                [embedding for _, _, embedding in modified_chunks],
                0,  # Start chunk index
                collection
            )

        content_hashes[file_name] = current_hash

# Monitor files for changes
def monitor_files(interval=10, max_unchanged_time=60):
    debug_print("Starting file monitor...")

    client = initialize_chromadb()
    if not client:
        raise RuntimeError("Failed to initialize ChromaDB client.")

    collection = client.get_or_create_collection("document_chunks")
    content_hashes = {}

    last_change_time = time.time()

    while True:
        current_time = time.time()
        if current_time - last_change_time >= max_unchanged_time:
            debug_print(f"No changes detected for {max_unchanged_time} seconds. Exiting monitor.")
            break

        process_files(content_hashes, collection)
        time.sleep(interval)

if __name__ == "__main__":
    try:
        monitor_files()
    except KeyboardInterrupt:
        debug_print("Monitor stopped by user")
    except Exception as e:
        debug_print(f"Error in monitor: {str(e)}")

"""
def validate_embeddings_and_storage():
    try:
        debug_print("Validating embeddings and storage in ChromaDB...")
        client = initialize_chromadb()
        if not client:
            raise RuntimeError("Failed to initialize ChromaDB client.")

        collection = client.get_or_create_collection("document_chunks")
        data = collection.get()

        if not data or not data.get("documents"):
            debug_print("No data found in ChromaDB. Ensure the processing code is working correctly.")
            return

        total_documents = len(data["documents"])
        debug_print(f"Total Chunks Stored: {total_documents}")

        debug_print("Sample Stored Data:")
        debug_print(f"Document: {data['documents'][0][:100]}...")  # Show first 100 characters
        debug_print(f"Metadata: {data['metadatas'][0]}")
        if "embeddings" in data:
            debug_print(f"Stored Embedding Lengths: {[len(emb) for emb in data['embeddings']]}")
        else:
            debug_print("No embeddings found.")
    except Exception as e:
        debug_print(f"Error validating embeddings and storage: {str(e)}")"""